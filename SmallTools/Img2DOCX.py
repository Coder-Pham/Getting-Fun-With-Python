'''Img2DOCX.py
This time I will do something different from 2 previous tools
I will use unlimited PI number to draw a IMG

Setting DOCX:
- Depend on ratio = width / height to Orientation
- Vertical:
    0.5cm in left/right margin
    0.5cm in top/bottom margin
    Width max = 190 characters in 5pt Arial
    Height max = 133 character in 5pt Arial no spacing
- Horizontal:
    0.5cm in left/right margin
    0.5cm in top/bottom margin
    Width max = 276 character in 5pt Arial
    Height max = 92 character in 5pt Arial no spacing

def hexcode(1 value)
    return '0x??'
    ?? from value in hex
'''

import time

from docx import Document
from docx.enum.section import WD_ORIENT  # Orientation for document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Alignment
from docx.shared import Pt, Cm  # Measurement in Word
# Pt for Line Spacing
# Cm for Page Margin
from docx.enum.text import WD_LINE_SPACING
from docx.shared import RGBColor    # RGB color formating
from mpmath import mp, pi
from PIL import Image


def hexcode(value):
    return '0x%02x' % (value)


def main():
    # Prepare for image
    File = input('Image file: ')
    count = time.time()

    img = Image.open(File)
    img = img.rotate(0, expand=1)
    width, height = img.size
    # DEBUG
    print('Width: {}, Height: {}, Ratio: {:0.5f}'.format(
        width, height, width/height))
    ratio = width / height
    # Reduce resolution
    if (ratio < 1):
        # Vertical
        resize_rate = height / 235
        height = int(height / resize_rate)
        width = int(ratio * height)
    else:
        # Horizontal
        resize_rate = width / 545
        if (int(width / (resize_rate * ratio)) > 164):
            resize_rate = height / 164
            width = int(ratio * height / resize_rate)
            height = 164
        else:
            width = int(width / resize_rate)
            height = int(width / ratio)
    img = img.resize((width, height))
    # DEBUG
    print('Width: {}, Height: {}, Ratio: {:0.5f}'.format(
        width, height, width/height))

    # Prepare for DOCX
    document = Document()
    # ORIENTATION
    section = document.sections[-1]
    if (ratio < 1):
        # Verical
        section.Orientation = WD_ORIENT.LANDSCAPE
    else:
        # Horizontal
        section.Orientation = WD_ORIENT.PORTRAIT

    # MARGIN
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)

    # PAGE TYPE
    section = document.sections[0]
    if (ratio < 1):
        # Vertical:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
    else:
        # Horizontal
        section.page_height = Cm(21)
        section.page_width = Cm(29.7)

    # Prepare for PI
    mp.dps = height * width
    pi = str(+mp.pi)
    pi = pi[:1] + pi[2:]

    # Process pixels
    for y in range(height):
        line = document.add_paragraph()
        for x in range(width):
            r, g, b = img.getpixel((x, y))
            char = line.add_run(pi[y*width + x])
            char.font.name = 'Arial'
            char.font.size = Pt(3)
            char.font.color.rgb = RGBColor(r, g, b)

        line.line_spacing_rule = WD_LINE_SPACING.SINGLE
        line.Alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = line.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = Pt(0)

    document.save('result.docx')
    print('Done with {:0.5f}s'.format(time.time() - count))


if __name__ == '__main__':
    main()
